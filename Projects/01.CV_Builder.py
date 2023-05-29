from docx import Document
from docx.shared import Inches
document=Document()

# Profile Picture
document.add_picture(
    'Projects/cat.jpg',
    width=Inches(2.0)
    )

# These lines of code are prompting the user to enter their name, phone number, and email address and
# storing the input values in the variables `name`, `PhNo`, and `email`, respectively.

name=input('Enter your name:')
PhNo=input('Enter your Phone number"')
email=input('Enter your Email ID:')

document.add_paragraph(name+'|'+str(PhNo)+'|'+email)

document.add_heading('About Me')

# about_me=input('Tell about your self:')

document.add_paragraph(input('Tell about your self:'))

# work experiance 
document.add_heading('Work Experience')


wrk_exp=document.add_paragraph()
company=input('Enter Company name:')
fromdate=input('From Date:')
todate=input('To Date:')

wrk_exp.add_run(company+' ').bold=True
wrk_exp.add_run(fromdate+'-'+todate+'\n').italic=True


#experiance 

exp_dtil=input('Describe your Experiance in '+company)

wrk_exp.add_run(exp_dtil)

# Loop work experiance 

while True:
    more_exp=input('Do you have more experiance? Yes / No :' )
    if more_exp.lower()=='yes':
                
        wrk_exp=document.add_paragraph()
        company=input('Enter Company name:')
        fromdate=input('From Date:')
        todate=input('To Date:')

        wrk_exp.add_run(company+' ').bold=True
        wrk_exp.add_run(fromdate+'-'+todate+'\n').italic=True


        #experiance 

        exp_dtil=input('Describe your Experiance in '+company)

        wrk_exp.add_run(exp_dtil)
    else:
        break
# Skills

document.add_heading('skills')

skill=input('Enter you skill set : ')
sk_pr=document.add_paragraph(skill)
sk_pr.style='List Bullet'

while True:
    more_skill=input('Do you have more skill? Yes / No :' )
    if more_skill.lower=='yes':
        skill=input('Enter you skill set : ')
        sk_pr=document.add_paragraph(skill)
        sk_pr.style='List Bullet'
    else:
        break

# Footer
FooterSec=document.sections[0]
footer=FooterSec.footer
Foot_Par=footer.paragraphs[0]
Foot_Par.text='CV Generated using Just Rise Technologies'

document.save('cv.docx')